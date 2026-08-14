# -*- coding: utf-8 -*-
"""生成《张睿-字节AI产品设计-简历.docx》—— 针对字节/抖音广告业务 AI 产品设计 JD 优化版。
不改动母版（张睿-标准简历数据库母版-可编辑版.docx），本脚本仅读取其内容用于核对，输出全新文件。
"""
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/rivizhang/Downloads/张睿-字节AI产品设计-简历.docx"

# 配色
NAVY = RGBColor(0x1F, 0x2A, 0x44)      # 标题深蓝
ACCENT = RGBColor(0x2F, 0x6D, 0xF0)   # 强调蓝（呼应字节）
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x22, 0x22, 0x22)

CN_FONT = "微软雅黑"

doc = Document()

# ---- 全局默认字体 / 页边距 ----
normal = doc.styles["Normal"]
normal.font.name = CN_FONT
normal.font.size = Pt(10)
normal.font.color.rgb = DARK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
rpr = normal._element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), CN_FONT)
rfonts.set(qn("w:hAnsi"), CN_FONT)
rfonts.set(qn("w:eastAsia"), CN_FONT)

sec = doc.sections[0]
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)


def set_run_font(run, size=None, bold=None, color=None, font=CN_FONT):
    run.font.name = font
    r = run._element.get_or_add_rPr()
    rf = r.get_or_add_rFonts()
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    rf.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color=ACCENT, size=6):
    p = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(bottom)
    p.append(pbdr)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=12.5, bold=True, color=NAVY)
    add_bottom_border(p, color=ACCENT, size=8)
    return p


def bullet(text_runs):
    """text_runs: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    # bullet symbol color
    for t, b in text_runs:
        r = p.add_run(t)
        set_run_font(r, size=10, bold=b, color=DARK if not b else NAVY)
    return p


def kv_line(label, value, label_color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10, bold=True, color=label_color)
    r2 = p.add_run(value)
    set_run_font(r2, size=10, color=DARK)
    return p


# ============ 头部 ============
name = doc.add_paragraph()
name.paragraph_format.space_after = Pt(1)
name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = name.add_run("张 睿")
set_run_font(r, size=22, bold=True, color=NAVY)

role = doc.add_paragraph()
role.paragraph_format.space_after = Pt(1)
r = role.add_run("求职意向：字节跳动 / 抖音集团 · 广告业务 AI 产品设计（实习生）")
set_run_font(r, size=11, bold=True, color=ACCENT)

contact = doc.add_paragraph()
contact.paragraph_format.space_after = Pt(1)
r = contact.add_run("北京  |  13067018267  |  1277929459@qq.com")
set_run_font(r, size=9.5, color=GRAY)

tags = doc.add_paragraph()
tags.paragraph_format.space_after = Pt(4)
r = tags.add_run("UI/UX 设计师 · AIGC 交互设计师 ·  AI Coding ·  模型评测 ·  Prompt 设计 ·  C 端 UX ·  Figma / AE ·  H5 原型")
set_run_font(r, size=9.5, color=GRAY)

# ============ 个人总结 ============
section_heading("个人总结")
summary = [
    ("清华美院科普信息设计硕士", True),
    ("（AI 辅助设计 / 多模态交互方向，GPA 3.9/4.0），兼具 AI 产品体验、UI/UX 与视觉设计复合能力。", False),
]
bullet(summary)
bullet([("深度参与 AI 原生产品的「模型 × 人」交互设计：", True),
        ("在阿里通义实验室参与万相文生图 / 图像编辑模型的体验与数据体系，搭建结构化 Prompt 体系与 Agent 评测，对模型能力边界、生成质量与人机交互有实践体感。", False)])
bullet([("具备真实 C 端产品从 0 到 1 / 改版上线经验", True),
        ("（腾讯新闻任务中心、字节抖音政务运营中台），覆盖需求拆解 → 信息架构 → 交互流程 → 高保真界面 → 动效 → 走查全链路，并以数据验证体验收益。", False)])
bullet([("高频使用 ChatGPT / Claude / Codex / Gemini / 即梦，能把 AI 用于需求分析、创意探索、交互验证与原型 Coding；", True),
        ("以第一作者发表 IEEE VR 2025 论文并入选 Gallery，持续思考「AI + 设计」的交互范式。", False)])

# ============ 教育背景 ============
section_heading("教育背景")
kv_line("清华大学美术学院", "科普信息设计 · 硕士（在读）    2024.09 – 2027.06")
bullet([("GPA 3.9/4.0；研究方向：用户体验设计、人工智能辅助设计、多模态交互与沉浸式体验设计。", False)])
kv_line("江南大学设计学院", "视觉传达设计 · 学士    2019.09 – 2023.06")
bullet([("优秀毕业设计；学业一等奖学金、校三好学生等荣誉。", False)])

# ============ 实习经历 ============
section_heading("实习经历")
# 字节（最相关，置顶）
kv_line("字节跳动 · 抖音政务运营中台", "视觉设计实习生    2025.07 – 2025.09")
bullet([("负责政务与文旅活动的视觉系统与 H5 体验设计，结合 AIGC 完成创意扩展与快速迭代，深入理解抖音产品矩阵与内容 / 营销玩法，推动 10+ 活动页面上线。", False)])
bullet([("参与活动 KV、互动 H5 与落地页全流程设计，沉淀可复用的 AIGC 创意模板，提升运营侧内容产出效率。", False)])
# 阿里通义
kv_line("阿里巴巴 · 通义实验室", "AI 视觉美学实习生    2025.10 – 2026.03")
bullet([("参与通义万相文生图 / 图像编辑模型的体验与数据体系优化，构建 SFT 多维数据标签并整理 8.3 亿张训练数据；设计「捏脸」「调色盘」等功能与结构化 Prompt 体系，推进模型能力向用户可用产品转化。", False)])
bullet([("参与万相 Agent 体验评测，完成 100+ 组样本测试并编写 2.6 版本 Prompt 指南，围绕生成稳定性、真实性与可控性推进体验优化，建立对 AI 能力边界的系统性认知。", False)])
# 腾讯新闻
kv_line("腾讯新闻", "产品体验设计实习生    2026.04 – 至今")
bullet([("任务中心改版：", True),
        ("负责信息架构、任务流程、积分反馈及移动端视觉升级，协同产品 / 研发推动 7/16 上线；改版后周日均页面进入 PV 由约 1.95 万提升至 2.27 万（增长约 13%），完成从概念到落地的完整设计交付。", False)])
bullet([("转盘抽奖：", True),
        ("设计「入口 → 抽奖 → 中奖反馈 → 记录 → 异常状态」完整交互链路；近 30 天累计抽奖 6,400+ 次，推动积分净消耗 9.26 万，验证可用性设计与转化效果。", False)])
bullet([("《攀了个岩》AI Coding 小游戏：", True),
        ("使用 Codex / Claude / WorkBuddy 完成玩法、UI 与 H5 原型开发并推动上线；连续 3 周在趣玩平台用户量、使用时长、留存量三项指标位列前三，近 14 天人均使用时长 204 秒、有效使用率 77.8%。", False)])

# ============ 项目经历 ============
section_heading("项目经历")
kv_line("北京亦庄模数世界 AIGC 展厅互动设计", "交互设计师")
bullet([("围绕 AIGC 生成内容与人机交互设计沉浸式展项，以脑电、面部表情等多模态数据驱动视觉生成；《情绪合成器》入选 IEEE VR 2025 Gallery 并以第一作者发表论文，沉淀对「原生 AI 产品交互范式」的思考。", False)])
kv_line("中国电信活动 AIGC 营销设计", "活动策划 / 产品体验设计师")
bullet([("参与广博会、元旦及春节线上活动，结合天翼云图 AIGC 完成蛇年 IP、3D 直播场景、H5 与抽奖机制设计，探索 AI 营销互动玩法，与广告 / 商业化业务高度契合。", False)])
kv_line("微信「状态」功能优化原型设计", "腾讯 WXG 产品经理训练营")
bullet([("围绕低频使用与表达场景单一问题开展用户访谈与需求拆解，完成信息架构、交互原型及 PRD，形成「用户研究 — 功能设计 — 原型验证」完整流程。", False)])

# ============ 获奖与学术 ============
section_heading("获奖与学术")
bullet([("清华大学梅贻琦奖学金｜IEEE VR 2025 第一作者论文 & Gallery 入选｜HRI 2025 LBR 论文收录｜「学院派奖」全国最高奖｜KTK 靳埭强设计奖优秀奖｜其他设计竞赛获奖 30+。", False)])

# ============ 专业技能 ============
section_heading("专业技能")
kv_line("产品体验", "用户研究、竞品分析、信息架构、交互流程、可用性分析、设计走查、PRD")
kv_line("视觉动效", "移动端 UI（iOS / Android 设计体系）· 组件规范 · 视觉系统 · 交互动效（After Effects）· 信息可视化 · 三维视觉")
kv_line("设计工具", "Figma（组件化协作）· Photoshop · Illustrator · After Effects · Premiere · Blender")
kv_line("AI 能力", "ChatGPT · Claude · Codex · Gemini · 即梦 · WorkBuddy；结构化 Prompt · 模型体验评测 · AI Coding · H5 交互原型")

doc.save(OUT)
print("SAVED:", OUT)
